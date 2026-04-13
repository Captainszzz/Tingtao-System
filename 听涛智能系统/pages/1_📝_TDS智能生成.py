import streamlit as st
from openai import OpenAI
import json
import re
import os
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# ================= 0. 🔑 在这里焊死你的 API Key =================
MY_API_KEY = st.secrets["MY_API_KEY"]

# ================= 1. Word 排版辅助函数 =================
def split_cn_en(text):
    if not text: return ""
    return re.sub(r'([\u4e00-\u9fa5]+)\s+([a-zA-Z].*)', r'\1\n\2', str(text))

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def format_cell_text(cell, text, is_header=False, is_property_col=False, align_left=False):
    cell.text = "" 
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if not cell.paragraphs: cell.add_paragraph()
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if align_left else WD_PARAGRAPH_ALIGNMENT.CENTER
    
    processed_text = str(text)
    if len(processed_text) < 12:
        processed_text = processed_text.replace("°C", "℃").replace(" C", " ℃")
        if processed_text.strip() == "C": processed_text = "℃"
    
    display_text = split_cn_en(processed_text) if is_property_col else processed_text
    run = paragraph.add_run(display_text)
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    if is_header:
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

def remove_row(table, row):
    table._tbl.remove(row._tr)

# ================= 2. 核心排版引擎 =================
def generate_docx(data):
    template_name = "TDS 模板.docx"
    cloud_path = f"听涛智能系统/{template_name}"
    
    if os.path.exists(cloud_path):
        doc = Document(cloud_path)
    elif os.path.exists(template_name):
        doc = Document(template_name)
    else:
        raise FileNotFoundError("找不到 Word 模板文件，请确保 'TDS 模板.docx' 和代码在同一个项目内！")

    processed_features = False
    processed_apps = False
    processed_props = False
    processed_guide = False

    for table in doc.tables:
        if len(table.rows) == 0: continue
        table_text = "".join([c.text for r in table.rows for c in r.cells])
        if any(kw in table_text for kw in ["声明", "Disclaimer", "备注"]): continue

        for i, row in enumerate(table.rows):
            cell_text = row.cells[0].text.strip().upper()
            if not processed_features and ("特性" in cell_text or "FEATURES" in cell_text) and "性能" not in cell_text:
                if i + 1 < len(table.rows):
                    items = data.get("Features", [])
                    content = "\n".join([f"• {x}" for x in items]) if items else "无"
                    format_cell_text(table.rows[i+1].cells[0], content, align_left=True)
                    processed_features = True
            
            if not processed_apps and ("推荐使用" in cell_text or "APPLICATIONS" in cell_text or "应用" in cell_text):
                if i + 1 < len(table.rows):
                    items = data.get("Applications", [])
                    content = "\n".join([f"• {x}" for x in items]) if items else "无"
                    format_cell_text(table.rows[i+1].cells[0], content, align_left=True)
                    processed_apps = True

        header_cell = table.rows[0].cells[0].text.strip()
        if not processed_props and any(kw in header_cell for kw in ["性能", "Properties"]) and len(table.rows[0].cells) >= 4:
            processed_props = True
            prop_data = {item.get("Property", "").strip(): item for item in data.get("TypicalProperties", [])}
            found_keys = set()
            to_delete = []

            for r in table.rows[1:]:
                tmp_name = r.cells[0].text.strip()
                match = None
                for k, v in prop_data.items():
                    if tmp_name in k or k in tmp_name:
                        match = v
                        found_keys.add(k)
                        break
                if match:
                    r.cells[0].text, r.cells[1].text = match.get("Property",""), match.get("TestMethod","")
                    r.cells[2].text, r.cells[3].text = match.get("Value",""), match.get("Unit","")
                else:
                    to_delete.append(r)
            
            for r in to_delete: remove_row(table, r)
            for k, v in prop_data.items():
                if k not in found_keys:
                    nr = table.add_row()
                    nr.cells[0].text, nr.cells[1].text = v.get("Property",""), v.get("TestMethod","")
                    nr.cells[2].text, nr.cells[3].text = v.get("Value",""), v.get("Unit","")

            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if i == 0:
                        set_cell_background(cell, "1F497D")
                        format_cell_text(cell, cell.text.strip(), is_header=True, is_property_col=False)
                    else:
                        format_cell_text(cell, cell.text.strip(), is_property_col=(j==0))
                        set_cell_background(cell, "F2F2F2" if i % 2 == 0 else "FFFFFF")

        elif not processed_guide and any(kw in header_cell for kw in ["加工", "Melt", "熔体"]):
            processed_guide = True
            for r in table.rows: remove_row(table, r)
            guide = data.get("ProcessingGuide", {})
            mapping = {
                "MeltTemp": "熔体温度 Melt Temp", "MoldTemp": "模温 Mold Temp",
                "BarrelZoneTemp": "料筒温度 Barrel Zone Temp", "InjectionSpeed": "注塑速度 Injection Speed",
                "BackPressure": "背压 Back Pressure", "DryingCondition": "干燥条件 Drying Condition",
                "ProcessingTemp": "加工上限温度 Processing Temp"
            }
            idx = 0
            for k, zh in mapping.items():
                val = guide.get(k, "")
                if val and val != "无":
                    nr = table.add_row()
                    for c_idx, cell in enumerate(nr.cells):
                        is_label = (c_idx == 0)
                        format_cell_text(cell, zh if is_label else val, is_property_col=False, align_left=is_label)
                        set_cell_background(cell, "F2F2F2" if idx % 2 == 0 else "FFFFFF")
                    idx += 1

    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()

# ================= 3. 网页界面与 AI 通信 =================
st.set_page_config(page_title="听涛 TDS 智能生成", page_icon="📝")
st.title("🚀 听涛新材料 - 全能 TDS 提取系统")
st.markdown("上传原厂资料，**支持图片、PDF、Word、Excel**，将自动阅读并生成 听涛TDS-Word文档。")

uploaded_file = st.file_uploader("📥 请上传原厂物料文件", type=['png', 'jpg', 'jpeg', 'pdf', 'docx', 'xlsx'])

if st.button("✨ 一键识别并生成 TDS", type="primary"):
    if not uploaded_file:
        st.error("❌ 请先上传文件哦！")
    else:
        try:
            with st.spinner(f"🧠 正在上传解析【{uploaded_file.name}】，阅读提取可能需要几十秒..."):
                
                client = OpenAI(
                    api_key=MY_API_KEY,
                    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
                )
                
                file_ext = os.path.splitext(uploaded_file.name)[1]
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
                    tmp_file.write(uploaded_file.getbuffer())
                    temp_file_path = tmp_file.name
                
                with open(temp_file_path, "rb") as f:
                    file_object = client.files.create(file=f, purpose="file-extract")

                prompt = """
                你是一个专业的数据提取员兼材料学翻译专家。请仔细阅读我上传的文件资料，提取其中的塑料物性表数据，并严格输出为以下的 JSON 格式。
                
                【💡 核心提取规则（非常重要）】：
                1. 必须且只能从上传的文件中提取数据，绝对不允许自己编造、猜测或联想任何数据！
                2. 对于 `TypicalProperties` 中的 `Property` 字段：必须强制输出【中文 + 空格 + 英文】的对照格式（例如："拉伸强度 Tensile Strength"）。
                3. 对于 `Features` (特性) 和 `Applications` (应用)：也必须强制输出中英对照格式。如果文件中没有提到特性或应用，请留空，不要编造！
                4. 不要回复任何其他说明文字，也不要包含 ```json 标签。
                
                输出格式模板：
                {
                  "ProductName": "提取物料名称",
                  "Features": ["优异的耐热性 Excellent heat resistance", "高流动性 High flowability"],
                  "Applications": ["汽车零部件 Automotive parts", "电子电器外壳 Electrical enclosures"],
                  "TypicalProperties": [
                    {"Property": "拉伸强度 Tensile Strength", "TestMethod": "ISO 527", "Value": "60", "Unit": "MPa"}
                  ],
                  "ProcessingGuide": {
                    "MeltTemp": "240-270 °C",
                    "MoldTemp": "60-80 °C"
                  }
                }
                """
                
                # =========================================================================
                # 【终极修复区】：阿里云专属通道，文件 ID 必须且只能放在 "system" 角色里！
                # =========================================================================
                response = client.chat.completions.create(
                    model="qwen-long", 
                    messages=[
                        {"role": "system", "content": f"fileid://{file_object.id}"},  # 👈 修复在这里：变成了 system
                        {"role": "system", "content": prompt},
                        {"role": "user", "content": "请严格按照 system 提示词的要求，读取上述内部系统文件，提取真实的物性表数据。绝不允许自己编造！"}
                    ]
                )
                
                result_text = response.choices[0].message.content
                
                clean_text = result_text
                if "```json" in clean_text:
                    clean_text = clean_text.split("```json")[1].split("```")[0].strip()
                elif "```" in clean_text:
                    clean_text = clean_text.split("```")[1].split("```")[0].strip()
                    
                try:
                    data = json.loads(clean_text)
                except Exception as e:
                    try: os.remove(temp_file_path) 
                    except: pass
                    st.error("❌ AI 返回的数据格式不标准，Word 生成失败！")
                    st.info(f"🕵️ 抓包看看 AI 到底说了什么：\n\n{result_text}")
                    st.stop()
                
                try: os.remove(temp_file_path)
                except: pass
                try: client.files.delete(file_object.id)
                except: pass
                
            with st.spinner("📝 数据提取成功！正在生成完美排版的文档..."):
                docx_bytes = generate_docx(data)
                
            st.success("🎉 听涛TDS生成啦！请点击下方按钮下载。")
            safe_name = str(data.get('ProductName', '新物料')).replace('/', '_')
            st.download_button(
                label="📥 下载 听涛TDS-Word文档，建议手动微调",
                data=docx_bytes,
                file_name=f"听涛TDS_{safe_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error(f"❌ 运行过程中出现错误：{e}")
            if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
                try: os.remove(temp_file_path)
                except: pass
